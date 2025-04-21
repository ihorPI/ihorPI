"""
BrassMouthpiece_demo.py - анализатор тембра тромбона и параметров мундштуков

Автор: [Пиданов]
Контакт: [figarokj@gmail.com]
Год: 2025

Лицензия: MIT License

Copyright (c) 2025 [Ваше имя]

Разрешается бесплатно использовать, копировать, изменять и распространять данный файл при соблюдении условий MIT License.
"""



import pyaudio
import PySimpleGUI as sg
import matplotlib
matplotlib.use('TkAgg')
import matplotlib.pyplot as plt
import matplotlib.figure as mpl_fig
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
import matplotlib.ticker as mticker
import matplotlib.cm as cm
import numpy as np
import wave
import time
import math
import sys
import os
import threading
import matplotlib.pyplot as plt
from scipy.io import wavfile
from scipy.signal import find_peaks
import warnings
from scipy.io.wavfile import WavFileWarning
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Название темы окна графического интерфейса
sg.theme ('Dark Blue 13')
# Определяем список файлов с расширением .wav
file_list = [f for f in os.listdir() if f.endswith('.wav')]
com_file_list = []


# Заданные значения частоты, соответствующим кнопкам
NOTES = ['Bb2', 'F3', 'Bb3', 'D4', 'F4', 'Ab4', 'Bb4']
VALUES = [117, 175, 233, 294, 349, 415, 466]
DIAPAZON = [5000, 10000, 10000, 15000, 15000, 20000, 20000]
button_frequency = 117  # или любое другое значение по умолчанию
button_num = 0
diapazon = DIAPAZON[button_num]  # Первое значение из DIAPAZON
button_values = NOTES[button_num]
# Начальное значение переменной trim в секундах
trim = 1
# Начальное значение переменной для шумов
noise_percentile = 90
tolerance_percentage = 0.25  # 25%
snr_threshold = 10  # Задаем запас в 10 дБ над шумом
num_harmonics = 40

# Путь логотипу


# ------ Начало кода графического интерфейса ----А-

# Параметры аудио
CHUNK = 1024
FORMAT = pyaudio.paInt16
CHANNELS = 1
RATE = 44100
THRESHOLD_DB = 6  # Порог активации записи (в дБ относительно шума)
NOISE_DURATION = 1  # Продолжительность измерения шума (в секундах)

# Инициализация переменных
is_recording = False
frames = []  # Список для хранения аудио данных по блокам

# -------- Начало кода Matplotlib --------
#
# Функция для отображения графика сигнала-1----Записи-----
def plot_signal(data, figure_canvas_agg, ax):
    
    ax.clear()
  
    ax.text(
            0.55, 1.0,  # Координаты (x, y) в относительных значениях
            'Форма звука ', 
            transform=ax.transAxes,  # Привязка координат графика (Axes)
            ha='right',          # Горизонтальное выравнивание текста
            va='bottom',         # Вертикальное выравнивание текста
            fontsize=10, 
            color='lightgray'
            )
 
 # Добавление сетки, которая адаптируется к шкале осей
    ax.grid(True)

    ax.set_xlim(0, CHUNK)
    #ax.set_ylim(-2000, 2000)
    ax.plot(data, color='lime', linewidth=0.5)



    figure_canvas_agg.draw()

# Функция для анализа спектра ----ОБЕРТОНЫ------ звука=2-----=
def plot_spectrum(data, figure_canvas_agg, ax, button_frequency, trim, noise_percentile, diapazon):
    ax.clear()
#___________вставить код__2______________________________начало_
# Шаг 03. Выполняем FFT для анализа частот ~~~~~~~~~~~~~>>>>>
        
    fft_spectrum = np.fft.fft(data)
    frequencies = np.fft.fftfreq(len(fft_spectrum), 1 / fs)
# Оставляем только положительные частоты
    positive_freqs = frequencies[:len(frequencies) // 2]
    positive_spectrum = np.abs(fft_spectrum[:len(frequencies) // 2])
# Ограничиваем спектр до 5000 Гц
    limit_freq = diapazon
    mask = positive_freqs <= limit_freq
    positive_freqs = positive_freqs[mask]
    positive_spectrum = positive_spectrum[mask]
# Шаг 04. Определение уровней шума по диапазонам частот
    def calculate_noise_levels(freqs, spectrum, num_bands):
        noise_levels = []
        band_edges = np.logspace(0, np.log10(limit_freq), num_bands + 1)  # Октавы или полосы
        for i in range(num_bands):
            band_mask = (freqs >= band_edges[i]) & (freqs < band_edges[i + 1])
            if np.any(band_mask):
                band_spectrum = spectrum[band_mask]
                noise_level = np.percentile(band_spectrum, noise_percentile)  # 90-й процентиль уровня шума
                noise_levels.append(20 * np.log10(noise_level))  # Перевод в дБ
            else:
                noise_levels.append(-np.inf)  # Если нет данных, присваиваем очень низкое значение
        return noise_levels, band_edges

    num_bands = diapazon // button_frequency  # Количество диапазонов
    noise_levels, band_edges = calculate_noise_levels(positive_freqs, positive_spectrum, num_bands)

# Шаг 05. Поиск пиков гармоник с учетом уровня шума
    peaks, _ = find_peaks(positive_spectrum)
    peak_frequencies = positive_freqs[peaks]
    peak_amplitudes = 20 * np.log10(positive_spectrum[peaks])

    
    tolerance = button_frequency * tolerance_percentage  # Расчет 15% от основной частоты

    harmonic_range = (button_frequency - tolerance,             button_frequency + tolerance)  # Диапазон поиска
    harmonic_mask = (peak_frequencies >= harmonic_range[0]) & (peak_frequencies <= harmonic_range[1])

    if np.any(harmonic_mask):
        harmonic_peak_idx = np.argmax(peak_amplitudes[harmonic_mask])
        button_frequency = peak_frequencies[harmonic_mask][harmonic_peak_idx]
        harmonic_amp = peak_amplitudes[harmonic_mask][harmonic_peak_idx]
    else:
        raise ValueError(f"Первая гармоника не найдена в диапазоне {harmonic_range[0]} - {harmonic_range[1]} Гц.")

# Шаг 07.1 Поиск следующих гармоник 40 гармоник
    
    harmonic_frequencies = [button_frequency]
    harmonic_amplitudes = [harmonic_amp]
    

    for n in range(2, num_harmonics + 1):
        target_freq = n * button_frequency
        harmonic_mask = (peak_frequencies >= target_freq - tolerance) & (peak_frequencies <= target_freq + tolerance)
    
        if np.any(harmonic_mask):
            harmonic_peak_idx = np.argmax(peak_amplitudes[harmonic_mask])
            harmonic_freq = peak_frequencies[harmonic_mask][harmonic_peak_idx]
            harmonic_amp = peak_amplitudes[harmonic_mask][harmonic_peak_idx]
        
 # Проверка уровня шума по диапазону
            band_idx = np.digitize(harmonic_freq, band_edges) - 1
            if band_idx >= 0 and band_idx < len(noise_levels):
# Добавляем порог SNR ( snr_threshold) в 6 дБ над уровнем шума
                if harmonic_amp > noise_levels[band_idx] + snr_threshold: 
# Для гармоники и амплитуды с учетом шума                
                    harmonic_frequencies.append(harmonic_freq)
                    harmonic_amplitudes.append(harmonic_amp)
                    
# Шаг 7.2 Поиск сильных частот между гармониками===============

# Для частот и амплитуд найденных сильных частот
    strong_frequencies = []
    strong_amplitudes = []
    strong_frequency_indices = []  # Для номеров найденных частот
    harmonic_indices = []  # Для номеров гармоник

    for n in range(1, len(harmonic_frequencies)):
# Определяем диапазон между двумя гармониками
        low_bound = harmonic_frequencies[n - 1] + tolerance
        high_bound = harmonic_frequencies[n] - tolerance

# Ищем пики в этом диапазоне
        range_mask = (peak_frequencies >= low_bound) & (peak_frequencies <= high_bound)
    
        if np.any(range_mask):
            strong_freqs = peak_frequencies[range_mask]
            strong_ampls = peak_amplitudes[range_mask]
        
 # Находим максимальную амплитуду в этом диапазоне
            max_idx = np.argmax(strong_ampls)
        
# Проверка уровня шума по диапазону
            band_idx = np.digitize(strong_freqs[max_idx], band_edges) - 1
            if band_idx >= 0 and band_idx < len(noise_levels) and strong_ampls[max_idx] > noise_levels[band_idx] + snr_threshold:  
                strong_frequencies.append(strong_freqs[max_idx])
                strong_amplitudes.append(strong_ampls[max_idx])
                strong_frequency_indices.append(max_idx)  # Сохраняем номер частоты
                harmonic_indices.append(n)  # Сохраняем номер гармоники============================
# Проверка условий между гармониками=================

# Шаг 07.2 Расчет центроида спектра============================
    centroid = np.sum(np.array(harmonic_frequencies) * np.array(harmonic_amplitudes)) / np.sum(harmonic_amplitudes)
    
    # Обновление текстового поля статус-бара
    window['-MEASURE-'].update(f"Центроид спектра___{centroid:.0f} Гц \nЧастота основного тона = {harmonic_frequencies[0]:.1f} Гц\n Амплитуда = {harmonic_amplitudes[0]:.0f} дБ", text_color='ivory')
    slopes_text = f"Нота: {button_values} , основной тон  {button_frequency:.2f} Гц"
    # Обновление текста в окне
    window['-SLOPES-'].update(slopes_text, text_color='ivory')
    
    

# Вывод результатов=========================

# Подсчет количества интермодуляционных частот
    num_intermodulation = len(strong_frequencies)
# Получение номеров (индексов) этих частот
    frequency_indices = [i for i, freq in enumerate(strong_frequencies)]
    num_noise_levels = len(noise_levels)
    
# Вывод ================================Текст------
# Шаг 08. Проверки и вывод текста

    # Количество обертонов 
    
    #num_harmonics = len(harmonic_frequencies)
    output_overtone = f"\n{len(harmonic_frequencies)} обертонов. "

# Влияние количества обертонов на тембр
    if len(harmonic_frequencies) >= 20:
        output_overtone += "Большое количество обертонов\n указывает на звонкость\n и пронзительность звука.\n"
    elif 14 <= len(harmonic_frequencies) < 20:
        output_overtone += "Умеренное количество обертонов\n создаёт  сбалансированный тембр,\n  более мягкий и чистый, без агрессивных\n или резких оттенков звук,\n ощущение плавности и гладкости тембра.\n"
    else:
        output_overtone += "Малое количество обертонов \nможет указывать на простой\n или тусклый звук.\n"
        # Добавляем output_text как метку через точку с прозрачным фоном (чтобы не отображалась на графике)
#    ax.scatter([], [], label=output_overtone, color='none') 

    
# ----------------- 2 гармоника ---------------
    output_text = ""
    if harmonic_amplitudes[1] >= harmonic_amplitudes[0] + 5 and harmonic_amplitudes[1] <= harmonic_amplitudes[0] + 10:
        output_text += "2-й обертон придает звуку резкости.\n"
    if harmonic_amplitudes[1] >= harmonic_amplitudes[0] + 3 and harmonic_amplitudes[1] <= harmonic_amplitudes[0] + 5:
        output_text += "2-й обертон придает звуку металлический оттенок.\n"
    if harmonic_amplitudes[1] >= harmonic_amplitudes[0] - 0 and harmonic_amplitudes[1] <= harmonic_amplitudes[0] + 3:
        output_text += "2-й обертон придает звуку округлость.\n"

    if harmonic_amplitudes[1] >= harmonic_amplitudes[0] - 6 and harmonic_amplitudes[1] <= harmonic_amplitudes[0] - 3:
        output_text += "2-й обертон придает звуку яркость.\n"
    if harmonic_amplitudes[1] >= harmonic_amplitudes[0] - 10 and harmonic_amplitudes[1] <= harmonic_amplitudes[0] - 6:
        output_text += " 2-й обертон придает звуку густоты.\n"
    if harmonic_amplitudes[1] >= harmonic_amplitudes[0] - 15 and harmonic_amplitudes[1] <= harmonic_amplitudes[0] - 10:
       output_text += " 2-й обертон придает звуку полноту.\n"

# ---------------- 3 гармоника ------------------
    if harmonic_amplitudes[2] >= harmonic_amplitudes[0] + 3 and harmonic_amplitudes[2] <= harmonic_amplitudes[0] + 6:
        output_text += "3-й обертон добавляет звуку плотности.\n"
    if harmonic_amplitudes[2] >= harmonic_amplitudes[0] - 0 and harmonic_amplitudes[2] <= harmonic_amplitudes[0] + 3:
        output_text += "3-й обертон добавляет звуку резкости.\n"
    if harmonic_amplitudes[2] >= harmonic_amplitudes[0] - 6 and harmonic_amplitudes[2] <= harmonic_amplitudes[0] - 0:
        output_text += "3-й обертон добавляет звуку металлические нюансы.\n"
    if harmonic_amplitudes[2] >= harmonic_amplitudes[0] - 10 and harmonic_amplitudes[2] <= harmonic_amplitudes[0] - 6:
        output_text += "3-й обертон добавляет звуку глубину.\n"
    if harmonic_amplitudes[2] >= harmonic_amplitudes[0] - 12 and harmonic_amplitudes[2] <= harmonic_amplitudes[0] - 10:
        output_text += "3-й обертон добавляет звуку полноту.\n"
    if harmonic_amplitudes[2] >= harmonic_amplitudes[0] - 20 and harmonic_amplitudes[2] <= harmonic_amplitudes[0] - 12:
        output_text += "3-й обертон добавляет звуку медности.\n"

# -------------------- 4 гармоника -------------
    if harmonic_amplitudes[3] >= harmonic_amplitudes[0] - 12 and harmonic_amplitudes[3] <= harmonic_amplitudes[0] - 10:
        output_text += "4-й обертон дает звуку глубокий оттенок.\n"
    if harmonic_amplitudes[3] >= harmonic_amplitudes[0] - 15 and harmonic_amplitudes[3] <= harmonic_amplitudes[0] - 12:
        output_text += "4-й обертон дает звуку округлости.\n"
    if harmonic_amplitudes[3] >= harmonic_amplitudes[0] - 25 and harmonic_amplitudes[3] <= harmonic_amplitudes[0] - 15:
        output_text += "4-й обертон звуку добавляет гармоничности.\n"

# --------------------- 5 гармоника ------------
    if harmonic_amplitudes[4] >= harmonic_amplitudes[0] - 16 and harmonic_amplitudes[4] <= harmonic_amplitudes[0] - 12:
        output_text += "5-й обертон усиливает объем.\n"
    if harmonic_amplitudes[4] >= harmonic_amplitudes[0] - 20 and harmonic_amplitudes[4] <= harmonic_amplitudes[0] - 16:
        output_text += "5-й обертон дает прозрачность.\n"
    if harmonic_amplitudes[4] >= harmonic_amplitudes[0] - 30 and harmonic_amplitudes[4] <= harmonic_amplitudes[0] - 20:
        output_text += "5-й обертон усиливает глубину.\n"
#++++++++++++++++++++++++++++++++++++++++++++++++++++++
#----ТЕМБР-------------------------------ТЕМБР-------
    base = harmonic_amplitudes[0]
    for i in range(1, 5):
        delta = harmonic_amplitudes[i] - base
        print(f"{i+1} гармоника: {harmonic_amplitudes[i]} дБ, разность с основной ({base} дБ): {delta:.1f} дБ")

   # ------------------------- ТЕМБР -----------------------
    output_timbre = "<> Тембральная окраска -  "
    if all([
    harmonic_amplitudes[1] >= harmonic_amplitudes[0] - 15 and harmonic_amplitudes[1] <= harmonic_amplitudes[0] - 5,
    harmonic_amplitudes[2] >= harmonic_amplitudes[0] -20 and harmonic_amplitudes[2] <= harmonic_amplitudes[0] - 10,
    harmonic_amplitudes[3] >= harmonic_amplitudes[0] - 25 and harmonic_amplitudes[3] <= harmonic_amplitudes[0] - 15,
    harmonic_amplitudes[4] >= harmonic_amplitudes[0] - 30 and harmonic_amplitudes[4] <= harmonic_amplitudes[0] - 20
]):
        output_timbre += "НАСЫЩЕННЫЙ ТЕМБР"


# ------------------------- ТЕМБР -----------------------
    if all([
    harmonic_amplitudes[1] >= harmonic_amplitudes[0] - 6 and harmonic_amplitudes[1] <= harmonic_amplitudes[0] - 0,
    harmonic_amplitudes[2] >= harmonic_amplitudes[0] - 12 and harmonic_amplitudes[2] <= harmonic_amplitudes[0] - 6,
    harmonic_amplitudes[3] >= harmonic_amplitudes[0] - 15 and harmonic_amplitudes[3] <= harmonic_amplitudes[0] - 12,
    harmonic_amplitudes[4] >= harmonic_amplitudes[0] - 20 and harmonic_amplitudes[4] <= harmonic_amplitudes[0] - 15
]):
        output_timbre += "КРИСТАЛЬНЫЙ ТЕМБР"

# ------------------------- ТЕМБР -----------------------
    if all([
    harmonic_amplitudes[1] >= harmonic_amplitudes[0] - 3 and harmonic_amplitudes[1] <= harmonic_amplitudes[0] - 0,
    harmonic_amplitudes[2] >= harmonic_amplitudes[0] - 6 and harmonic_amplitudes[2] <= harmonic_amplitudes[0] - 3,
    harmonic_amplitudes[3] >= harmonic_amplitudes[0] - 12 and harmonic_amplitudes[3] <= harmonic_amplitudes[0] - 6,
    harmonic_amplitudes[4] >= harmonic_amplitudes[0] - 20 and harmonic_amplitudes[4] <= harmonic_amplitudes[0] - 12
]):
        output_timbre += "ПЛОТНЫЙ ТЕМБР"

# ------------------------- ТЕМБР -----------------------
    if all([
    harmonic_amplitudes[1] >= harmonic_amplitudes[0] - 3 and harmonic_amplitudes[1] <= harmonic_amplitudes[0] - 0,
    harmonic_amplitudes[2] >= harmonic_amplitudes[0] - 6 and harmonic_amplitudes[2] <= harmonic_amplitudes[0] - 3,
    harmonic_amplitudes[3] >= harmonic_amplitudes[0] - 15 and harmonic_amplitudes[3] <= harmonic_amplitudes[0] - 10,
    harmonic_amplitudes[4] >= harmonic_amplitudes[0] - 20 and harmonic_amplitudes[4] <= harmonic_amplitudes[0] - 15
]):
        output_timbre += "РЕЗКИЙ ТЕМБР"

# ------------------------- ТЕМБР -----------------------
    if all([
    harmonic_amplitudes[1] >= harmonic_amplitudes[0] - 3 and harmonic_amplitudes[1] <= harmonic_amplitudes[0] + 3,
    harmonic_amplitudes[2] >= harmonic_amplitudes[0] - 10 and harmonic_amplitudes[2] <= harmonic_amplitudes[0] - 6,
    harmonic_amplitudes[3] >= harmonic_amplitudes[0] - 20 and harmonic_amplitudes[3] <= harmonic_amplitudes[0] - 15,
    harmonic_amplitudes[4] >= harmonic_amplitudes[0] - 20 and harmonic_amplitudes[4] <= harmonic_amplitudes[0] - 15
]):
        output_timbre += "ЯРКИЙ ТЕМБР"

# ------------------------- ТЕМБР -----------------------
    if all([
    harmonic_amplitudes[1] >= harmonic_amplitudes[0] - 6 and harmonic_amplitudes[1] <= harmonic_amplitudes[0] - 3,
    harmonic_amplitudes[2] >= harmonic_amplitudes[0] - 12 and harmonic_amplitudes[2] <= harmonic_amplitudes[0] - 10,
    harmonic_amplitudes[3] >= harmonic_amplitudes[0] - 15 and harmonic_amplitudes[3] <= harmonic_amplitudes[0] - 10,
    harmonic_amplitudes[4] >= harmonic_amplitudes[0] - 20 and harmonic_amplitudes[4] <= harmonic_amplitudes[0] - 15
]):
        output_timbre += "ТЕПЛЫЙ ТЕМБР"
    else:
        output_timbre += "- неповторимый обертоновый рисунок <>"
        # Добавляем output_odd как метку через точку с прозрачным фоном (чтобы не отображалась на графике)
#    ax.scatter([], [], label=output_timbre, color='none') 

        
       
#----ТЕМБР-------------------------------ТЕМБР-------
#++++++++++++++++++++++++++++++++++++++++++++++++++++    
# Сравнение суммы амплитуд нечетных и четных гармоник
    # Получаем нечетные и четные гармоники
    odd_list = harmonic_amplitudes[::2]
    even_list = harmonic_amplitudes[1::2]

# Ограничиваем обе группы до одинаковой длины
    min_len = min(len(odd_list), len(even_list))
    odd_harmonics = sum(odd_list[:min_len])
    even_harmonics = sum(even_list[:min_len])
    # Добавляем анализ по соотношению гармоник
    output_odd = ""
    if odd_harmonics > even_harmonics:
        output_odd += "\nНечётные обертоны преобладают, что придаёт звуку\n матовый, более приглушённый оттенок.\n"
    else:
        output_odd += "\nЧётные обертоны преобладают, что делает звук\n более ярким и насыщенным.\n"
        
# Печать суммы амплитуд в нечетных и четных гармоник
    
    print(f"Сумма нечетных гармоник: {odd_harmonics:.0f}\n")

    print(f"Сумма четных гармоник: {even_harmonics:.0f}")

        
    # Добавляем output_odd как метку через точку с прозрачным фоном (чтобы не отображалась на графике)
#    ax.scatter([], [], label=output_odd, color='none') 
        
     # Суммарная амплитуда по диапазонам
    low_range = sum(harmonic_amplitudes[1:5])  # 2–5 обертона ~200–700 Гц
    mid_range = sum(harmonic_amplitudes[5:15])  # 6–15 обертонов
    high_range = sum(harmonic_amplitudes[15:])  # 16+ обертонов

    output_diapazon = ""
    if low_range > mid_range and low_range > high_range:
        output_diapazon += "\nВыражены нижние гармоники\n — звук глубокий и насыщенный.\n"
    if mid_range > low_range and mid_range > high_range:
        output_diapazon += "\nСерединные гармоники доминируют\n — звук яркий и звонкий.\n"
    if high_range > mid_range and high_range > low_range:
        output_diapazon += "\nВысокие гармоники преобладают\n — звук резкий или пронзительный.\n"
# Добавляем output_diapazon как метку через точку с прозрачным фоном (чтобы не отображалась на графике)
#    ax.scatter([], [], label=output_diapazon, color='none')
        
        

    # Проверка пиков в характерных областях
    brightness_band = [i for i, f in enumerate(harmonic_frequencies) if 2500 <= f <= 3000]
    brightness_level = sum(harmonic_amplitudes[i] for i in brightness_band)
    
    # Печать суммы амплитуд в этом диапазоне
    print(f"Сумма амплитуд в диапазоне 2.5–3 кГц: {brightness_level:.0f}")

    output_spector = ""
    if brightness_level > 0:
        output_spector += "\nНаблюдается пик в диапазоне 2.5–3 кГц\n — звук приобретает полетность и звонкость.\n"

    harshness_band = [i for i, f in enumerate(harmonic_frequencies) if 3000 <= f <= 4500]
    harshness_level = sum(harmonic_amplitudes[i] for i in harshness_band)
    
    # Печать суммы амплитуд в этом диапазоне
    print(f"Сумма амплитуд в диапазоне 3–4.5 кГц: {harshness_level:.0f}")
    if harshness_level > 0:
        output_spector += "\nВ диапазоне 3–4.5 кГц активность высоких гармоник\n — звук может быть пронзительным или резким.\n"

       
        

            
    output_centroid = f"CF = {centroid:.0f} Гц — "
    if centroid < 1500:
        output_centroid += "звук воспринимается\n как глубокий и тёмный.\n"
    elif centroid < 3000:
        output_centroid += "звук сбалансирован,\n умеренно яркий.\n"
    else:
        output_centroid += "звук яркий.\n"
    
    
    #return output_centroid   
    
     # Добавляем output_centroid как метку через точку с прозрачным фоном (чтобы не отображалась на графике)
#    ax.scatter([], [], label=output_centroid, color='none')  
    
# Вывод ========================

# Визуализация спектра >>>>>>>>>>>>>>>>>>>>>>>>>
    ax.plot(positive_freqs, 20 * np.log10(positive_spectrum), color='deepskyblue', linewidth=0.1)
# Вывод ========================
# Используем ax.scatter для отображения обертонов
    ax.scatter( harmonic_frequencies, harmonic_amplitudes, color='red', 
label=f'Обертоны: {len(harmonic_frequencies)} шт', 
s=10, 
zorder=3)
    
# Используем ax.scatter для отображения интермодуляционных частот
    ax.scatter(
    strong_frequencies, 
    strong_amplitudes, 
    color='#FFDC00', 
    label=f'Интермодуляционные частоты: {num_intermodulation} шт\nНомера частот: {harmonic_indices}', 
    s=7, 
    zorder=5
)
    output_mod = f'{num_intermodulation} - интермодуляционных частот: {harmonic_indices}'
# Вывод ========================

# Ограничиваемся только первыми 10 сильными частотами
    num_strong_frequencies = min(10, len(strong_frequencies))
    output_intermodul = ""
    for i in range(num_strong_frequencies):
        if i < len(harmonic_amplitudes) and i <     len(strong_amplitudes) and i < len(harmonic_indices):
        # Определяем разницу
            harmonic_amp = harmonic_amplitudes[i]
            intermodulation_amp = strong_amplitudes[i]
            difference = harmonic_amp - intermodulation_amp

        # Инициализируем influence_label значением по умолчанию
            influence_label = "-влияние не ощущается,  \n"

        # Условия для уровней влияния с учетом индекса гармоники
            if i == 0:  # 1-я гармоника
                threshold_min, threshold_max = 15, 20
                if difference < threshold_min:  
                    influence_label = "-придает мягкость."
                elif threshold_min <= difference <= threshold_max:
                    influence_label = "-усиливает мягкие оттенки."

            elif i == 1:  # 2-я гармоника
                threshold_min, threshold_max = 10, 20
                if difference < threshold_min:
                    influence_label = "-придает теплоты."
                elif threshold_min <= difference <= threshold_max:
                    influence_label = "-усиливает теплые оттенки."

            elif i == 2:  # 3-я гармоника
                threshold_min, threshold_max = 10, 15
                if difference < threshold_min:
                    influence_label = "-придает глубины."
                elif threshold_min <= difference <= threshold_max:
                    influence_label = "-усиливает глубину звука."

            elif i == 3:  # 4-я гармоника
                threshold_min, threshold_max = 12, 15
                if difference < threshold_min:
                    influence_label = "-придает текстурность."
                elif threshold_min <= difference <= threshold_max:
                    influence_label = "-усиливает текстуру звука."

            elif i == 4:  # 5-я гармоника
                threshold_min, threshold_max = 6, 12
                if difference < threshold_min:
                    influence_label = "-Пространственное звучание."
                elif threshold_min <= difference <= threshold_max:
                    influence_label = "-Шероховатость."

            elif i == 5:  # 6-я гармоника
                threshold_min, threshold_max = 6, 12
                if difference < threshold_min:
                    influence_label = "-Шум."
                elif threshold_min <= difference <= threshold_max:
                    influence_label = "-Пришептывание."

            elif i >= 6:  # 7-я и выше гармоники
                threshold_min, threshold_max = 15, 20
                if difference < threshold_min:
                    influence_label = "-Призвуки."
                elif threshold_min <= difference <= threshold_max:
                    influence_label = "-Небольшие призвуки."

# Формируем метку с частотой, амплитудой и результатом сравнения
            output_intermodul += (f"F({harmonic_indices[i]})={strong_frequencies[i]:.0f} Гц  {influence_label}\n")
#            output_mod += label
    # Отображаем точку на графике с уникальной меткой для каждой точки сильных частот
#            ax.scatter(strong_frequencies[i], strong_amplitudes[i], color='#FFDC00', s=2, label=label)
            


#=================================================
# Добавляем линию уровня шума только для десятого уровня
    ax.axhline(
    y=noise_levels[-1], 
    color='red', 
    linestyle='-', 
    linewidth=0.7,
    label=f'Уровень шума {noise_levels[-1]:.0f} дБ'
)
    print(f'Уровень шума 1 _:{noise_levels[-1]:.0f} дБ')
#==============================================
# Используем ax.axvspan для отображения диапазонов шума
    for i in range(num_bands):
        ax.axvspan(
        band_edges[i], 
        band_edges[i + 1], 
        color='gray',  # Светло-синий
        alpha=0.1, 
        label='' if i == 0 else ""
    )

# Включаем легенду, чтобы отобразить метку
    # Устанавливаем параметры легенды
    legend = ax.legend(loc='upper right', fontsize=7, bbox_to_anchor=(1.06, 1.05))
    legend.get_frame().set_facecolor((0, 0.2, 0.4, 0.1))  # Полупрозрачный черный фон
    #legend.get_frame().set_alpha(0)  # Полностью прозрачный фон
    legend.get_frame().set_edgecolor('none')  # Убираем границу
    for text in legend.get_texts():
        text.set_color("lightgray")  # Светло-серый текст
#==============================================

# Устанавливаем заголовок и метки осей, а также другие параметры через ax
    #ax.set_title(f'Спектр звука: {selected_file} и его анализ')
    ax.text(
    0.6, 1.0, # Координаты (x, y) в относительных значениях
    f'Анализ спектра: {selected_file}          ', 
    transform=ax.transAxes,  # Привязка координат графика (Axes)
    ha='right',          # Горизонтальное выравнивание текста
    va='bottom',         # Вертикальное выравнивание текста
    fontsize=10, 
    color='lightgray'
)
    
    #ax.set_xlabel('Частота (Гц)')
    #ax.set_ylabel('Амплитуда (дБ)')
    #текст по осям внутри графика------
    ax.text(0.95, 0.01, 'Frequency, Hz',
       verticalalignment='bottom', horizontalalignment='right',
       transform=ax.transAxes,
        color='ivory', fontsize=6)

    #ax.text(0.02, 0.8, 'Amplitude, dB', fontsize=6, rotation='vertical', ha='center', va='baseline', color='ivory')
    
    
    ax.set_xlim(0, diapazon)
    #ax.legend()  # Отображаем легенду
    ax.grid(True)  # Включаем сетку
#_вставить===
    # Позиция для аннотаций
    y_pos = -0.12  # Позиция по Y ниже графика
    x_step = 0.11  # Шаг по X для каждого блока

# Добавление блоков аннотаций
    for i in range(10):
    # Определяем координаты аннотации
        x_pos = i * x_step

    # Текст аннотации
        annotation_text = (f"{harmonic_frequencies[i]:.0f} Hz\n"
                       f"{harmonic_amplitudes[i]:.0f} dB\n"
                       f"  {i+1}")
        

    # Добавляем аннотацию на график
        ax.text(x_pos, y_pos, annotation_text,
            ha='center', va='top', transform=ax.transAxes,
            color="ivory", fontsize=6, bbox=dict(facecolor="black", edgecolor="black", boxstyle="square,pad=1.5", alpha=0.1))


#_вставить===========================================================0
# Добавление блок печати:
        
#    for i in range(len(harmonic_frequencies)):
      
#       print(f"{i+1} обертон "f"{harmonic_frequencies[i]:.0f} Hz  "
#                       f"{harmonic_amplitudes[i]:.0f} dB "  )
       
#=========================== Огибающая от 2,5 до 4 кГц========== 
    # Разделение на диапазоны
    base_band = [i for i, f in enumerate(harmonic_frequencies) if f < 2500]
    brightness_band = [i for i, f in enumerate(harmonic_frequencies) if 2500 <= f <= 3000]
    harshness_band = [i for i, f in enumerate(harmonic_frequencies) if 3000 <= f <= 4500]

# Вычисление средних амплитуд
    base_avg = np.mean([harmonic_amplitudes[i] for i in base_band]) if base_band else 0
    brightness_avg = np.mean([harmonic_amplitudes[i] for i in brightness_band]) if brightness_band else 0
    harshness_avg = np.mean([harmonic_amplitudes[i] for i in harshness_band]) if harshness_band else 0

# Вывод средних значений
    print(f"\n🎚️ Средние значения амплитуд:")
    print(f"До 2.5 кГц (гармоники {[i+1 for i in base_band]}): {base_avg:.0f} дБ")
    print(f"2.5–3 кГц (гармоники {[i+1 for i in brightness_band]}): {brightness_avg:.0f} дБ")
    print(f"3–4.5 кГц (гармоники {[i+1 for i in harshness_band]}): {harshness_avg:.0f} дБ")

# Анализ диапазона 2.5–3 кГц
    diff_brightness = brightness_avg - base_avg
    if diff_brightness <= -30:
        bright_text = " Большой спад — звонкость не ощущается, ближе к мягкому звучанию."
    elif -30 < diff_brightness <= -25:
        bright_text = " Спад — лёгкий намёк на звонкость."
    elif -25 < diff_brightness <= -20:
        bright_text = " Небольшой спад — умеренная, неагрессивная звонкость придающая полетность."
    elif -20 < diff_brightness <= -15:
        bright_text = " Подъем — яркий и энергичный звук, хорошо проецируется."
    elif diff_brightness > -15:
        bright_text = " Резкий подъем — отчётливо выраженная звонкость."
    else:
        
        bright_text = " Неопределённая характеристика звонкости."

# Анализ диапазона 3–4.5 кГц
    diff_harsh = harshness_avg - brightness_avg
    if diff_harsh <= -20:
        harsh_text = " Пассивность высоких гармоник – упрощает звук, делает звук более чистым, без агрессивных или резких оттенков."
    elif -20 < diff_harsh <= -10:
        harsh_text = " Умеренность высоких гармоник — звук приобретает сочность."
    elif -10 < diff_harsh <= 0:
        harsh_text = " Активность высоких гармоник — возможна пронзительность звука."
    else:
        harsh_text = " Неопределённая характеристика верхних гармоник."

# Печать результата
#    print("\n🎼 Заключение:")
#    print(f"• В диапазоне 2.5–3 кГц наблюдается: {bright_text}")
#    print(f"• В диапазоне 3–4.5 кГц наблюдается: {harsh_text}")

#============================    
    # === СОЗДАНИЕ ДОКУМЕНТА ===
    doc = Document()
    # Установка метаданных
    core_props = doc.core_properties
    core_props.author = "Пиданов"
    core_props.title = "BrassMouthpiece Demo"
    core_props.subject = "Анализ тембра тромбона и мундштуков"
    core_props.comments = "Бесплатно. Используйте с упоминанием автора."
    core_props.keywords = "Brass, Trombone, Mouthpiece, Тембр, Анализ, Пиданов"
    
    # Верхний колонтитул: логотип
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.text = "Автор: Пиданов, 2025 | Бесплатно. Используйте с упоминанием автора | CC BY-NC-SA"

# (опционально) Выравнивание по центру или справа:
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # или CENTER, LEFT
    
# Заголовок=========================================================
    doc.add_heading("Анализатор обертонов (тембра) тромбона и параметров мундштуков", level=1)
    # Добавляем абзац с нижней границей (линией)
    p = doc.add_paragraph()
    p_format = p.paragraph_format

# Добавление нижней границы через xml
    p_borders = OxmlElement('w:pBdr')
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single')
    bottom_border.set(qn('w:sz'), '6')       # Толщина
    bottom_border.set(qn('w:space'), '1')
    bottom_border.set(qn('w:color'), 'auto')
    p_borders.append(bottom_border)
    p._p.get_or_add_pPr().append(p_borders)
    
    
    doc.add_paragraph(f"Анализ проведён: {datetime.today().strftime('%d.%m.%Y')}\nПриложение BrassMouthpiece.exe ", style='Normal')
        
# === СОХРАНЕНИЕ И ВСТАВКА ИЗОБРАЖЕНИЯ ===    
    # Получаем имя файла без расширения
    base_filename = os.path.splitext(selected_file)[0]
# Сохраняем график
    image_path = f"Спектр_{base_filename}.png"
    fig.savefig(image_path)    
# Вставляем изображение по центру
    doc.add_paragraph(f"График спектра звука аудио файла:  {selected_file}", style='Heading 2')
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(image_path, width=Inches(7))
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
#-----------------тон--------
#    para = doc.add_paragraph()
    doc.add_paragraph(' '.join(line.strip() for line in slopes_text.splitlines() if line.strip()))


# Раздел 1: Тембровый профиль
    doc.add_heading("Тембровый профиль инструмента", level=2)
   

# Первый блок: объединяем в один абзац,---------------1
    
    para = doc.add_paragraph()
    para.add_run("1. Обертоны ").bold = True
    para.add_run("– это составляющие звукового спектра, кратные частоте основного тона (f₀). Их количество, амплитуда, расположение и соотношение определяют характер звучания инструмента.")
    para = doc.add_paragraph()
    para.add_run("•").bold = True
    para.add_run("Обертоны:").bold = True
    # Устанавливаем интервал после абзаца в 0
    para.paragraph_format.space_after = 0
# Последний блок: +

# Второй текст: 
    doc.add_paragraph(' '.join(line.strip() for line in output_overtone.splitlines() if line.strip()))


# Второй блок: опять нормальный абзац ----------------2
    
    para = doc.add_paragraph()
    para.add_run("•").bold = True
    para.add_run("Первые (1–5) гармоники  (обертоны) ").bold = True
    para.add_run(" обладают наибольшей амплитудой  и определяют фундаментальные  свойства тембра, такие как его теплота  и глубина. В данном контексте соотношение уровней обертонов (2–5) к основному тону вносит дополнительные оттенки в тембр:")
    # Устанавливаем интервал после абзаца в 0
    para.paragraph_format.space_after = 0
# Последний текст 
    for line in output_text.splitlines():
        line = line.strip()
        if line:
            doc.add_paragraph(f"– {line}")


# ТРЕТИЙ блок: опять нормальный абзац------------------3
    para = doc.add_paragraph()
    para.add_run("•").bold = True
    para.add_run("Отношение четных обертонов к нечетным:").bold = True
# Последний блок: +
    # Устанавливаем интервал после абзаца в 0
    para.paragraph_format.space_after = 0
    doc.add_paragraph(' '.join(line.strip() for line in output_odd.splitlines() if line.strip()))

# 4 блок: опять нормальный абзац -----------------------4
    para = doc.add_paragraph()
    para.add_run("•").bold = True
    para.add_run("Распределение энергии обертонов по диапазону:").bold = True
    # Устанавливаем интервал после абзаца в 0
    para.paragraph_format.space_after = 0
# Последний блок: +
    doc.add_paragraph(' '.join(line.strip() for line in output_diapazon.splitlines() if line.strip()))
    

    # 5-6 блок: опять нормальный абзац ----------------------6
  
    para = doc.add_paragraph()
    para.add_run("2. Акустически, огибающая спектра ").bold = True   
    para.add_run("описывает общую форму распределения энергии. \n")
    para.add_run("•").bold = True
    para.add_run("В диапазоне 2.5–3 кГц наблюдается:").bold = True
    # Устанавливаем интервал после абзаца в 0
    para.paragraph_format.space_after = 0
# Последний блок: нумерованный список
    doc.add_paragraph(' '.join(line.strip() for line in bright_text.splitlines() if line.strip()))
    
    # 7 блок: опять нормальный абзац ----------------------7

    para = doc.add_paragraph()
    para.add_run("•").bold = True
    para.add_run("В диапазоне 3–4.5 кГц наблюдается:").bold = True
    # Устанавливаем интервал после абзаца в 0
    para.paragraph_format.space_after = 0
# Последний блок: нумерованный список
    doc.add_paragraph(' '.join(line.strip() for line in harsh_text.splitlines() if line.strip()))
    
    # 8 блок: опять нормальный абзац ----------------------8
    
    para = doc.add_paragraph()
    para.add_run("•").bold = True
    para.add_run("Центроид спектра ").bold = True
    para.add_run(" - это характеристика звукового спектра, которая позволяет определить центр тяжести распределения обертонов, отвечает за яркость:")
    # Устанавливаем интервал после абзаца в 0
    para.paragraph_format.space_after = 0
# Последний блок: нумерованный список
    doc.add_paragraph(' '.join(line.strip() for line in output_centroid.splitlines() if line.strip()))
    
    # 9 блок: опять нормальный абзац ----------------------8
# Разбиваем на части для форматирования
    para = doc.add_paragraph()
    para.add_run("3. Психоакустические аспекты\n").bold = True   
    para.add_run("Интермодуляционные частоты - это частоты, которые возникают в результате взаимодействия обертонов и основного тона. "
             "Они могут влиять на восприятие тембра.\n")
    para.add_run("•").bold = True
    para.add_run("Меньшее количество интермодуляций делает звук более чистым, прозрачным и легко воспринимаемым:").bold = True
    # Устанавливаем интервал после абзаца в 0
    para.paragraph_format.space_after = 0
             

# Последний блок: нумерованный список (пример, без изменения)
    doc.add_paragraph(' '.join(line.strip() for line in output_mod.splitlines() if line.strip()))
    
    # 10 блок: опять нормальный абзац ----------------------10
    para = doc.add_paragraph()
    para.add_run("•").bold = True
    para.add_run("При большем числе интермодуляций тембр становится сложнее — появляются шероховатость, насыщенность и дополнительные оттенки:").bold = True 
    # Устанавливаем интервал после абзаца в 0
    para.paragraph_format.space_after = 0  

# Последний блок: нумерованный список
    doc.add_paragraph(' '.join(line.strip() for line in output_intermodul.splitlines() if line.strip()))
    
# ТЕМБР блок: опять нормальный абзац -----ТЕМБР--------------Т
    # Добавляем абзац с нижней границей (линией)
    p = doc.add_paragraph()
    p_format = p.paragraph_format

# Добавление нижней границы через xml
    p_borders = OxmlElement('w:pBdr')
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single')
    bottom_border.set(qn('w:sz'), '6')       # Толщина
    bottom_border.set(qn('w:space'), '1')
    bottom_border.set(qn('w:color'), 'auto')
    p_borders.append(bottom_border)
    p._p.get_or_add_pPr().append(p_borders)
    # Создаём абзац
    paragraph = doc.add_paragraph()

# Добавляем run со всем текстом
    run = paragraph.add_run(' '.join(line.strip() for line in output_timbre.splitlines() if line.strip()))

# Применяем тёмно-синий цвет
    run.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue
    run.bold = True
    run.italic = True
    
#=====================================================================0
# === СОХРАНЕНИЕ И ВСТАВКА ИЗОБРАЖЕНИЯ ===


    # === СОХРАНЕНИЕ ДОКУМЕНТА ===

    doc_filename = f"Анализ_{base_filename}.docx"
    doc.save(doc_filename)
    print(f"Файл успешно сохранён: {doc_filename}")
    update_status_bar(window, f"Файл успешно сохранён:\n{doc_filename}", text_color='#FFA500')
#_вставить код__2______________________________конец__
    figure_canvas_agg.draw()

# Функция внизу, для сравнения спектров трех файлов-3----------=

    
    # --- Функции для обработки данных ---
def extract_harmonics(file, trim, limit_freq, num_harmonics, button_frequency, tolerance_percentage):
    try:
        fs, data = wavfile.read(file)
    except Exception as e:
        raise ValueError(f"Ошибка чтения файла {file}: {e}")
    
    if data.ndim > 1:
        data = data[:, 0]
    data = data[:int(trim * fs)]
    
    fft_spectrum = np.fft.fft(data)
    frequencies = np.fft.fftfreq(len(fft_spectrum), 1 / fs)
    positive_freqs = frequencies[:len(frequencies) // 2]
    positive_spectrum = np.abs(fft_spectrum[:len(frequencies) // 2])
    
    mask = positive_freqs <= limit_freq
    positive_freqs = positive_freqs[mask]
    positive_spectrum = positive_spectrum[mask]
    
    peaks, _ = find_peaks(positive_spectrum)
    peak_frequencies = positive_freqs[peaks]
    peak_amplitudes = 20 * np.log10(positive_spectrum[peaks])
    
    harmonics_freqs = []
    harmonics_amps = []
    tolerance = button_frequency * tolerance_percentage
    results = []
    
    for n in range(1, num_harmonics + 1):
        target_freq = n * button_frequency
        harmonic_mask = (peak_frequencies >= target_freq - tolerance) & (peak_frequencies <= target_freq + tolerance)
        if np.any(harmonic_mask):
            idx = np.argmax(peak_amplitudes[harmonic_mask])
            harmonics_freqs.append(peak_frequencies[harmonic_mask][idx])
            harmonics_amps.append(peak_amplitudes[harmonic_mask][idx])
        else:
            break
        
        
    
    
    if harmonics_amps:
        max_amp = max(harmonics_amps)
        harmonics_amps = [amp / max_amp for amp in harmonics_amps]
    
    return harmonics_freqs, harmonics_amps  # Не ограничиваем количество гармоник для графика в интерфейсе -3


def compare_spectrums(harmonics_data, selected_files, figure_canvas_agg, ax):
    """
    Функция для сравнения спектров гармоник нескольких файлов.
    
    Параметры:
    harmonics_data -- данные гармоник (частоты и амплитуды) для каждого файла
    selected_files -- список имён файлов
    figure_canvas_agg -- объект FigureCanvasTkAgg для отрисовки
    ax -- объект оси графика (Axes) для отрисовки
    """
    # Очистка предыдущего содержимого графика
    ax.clear()

    # Установка параметров графика
    colors = ['blue', 'orange', 'red', 'purple', 'green' ]
    markers = ['.', '.', '.', '.', '.']
    centroids = []  # Хранение центроидов
    slopes = []  # Хранение коэффициентов наклона

    for i, (harmonic_freqs, harmonic_amps) in enumerate(harmonics_data):
    # Расчет центроида только для первых 20 гармоник
        limited_freqs = harmonic_freqs[:20]
        limited_amps = harmonic_amps[:20]

        if len(limited_amps) > 0:  # Проверяем, что ограниченные амплитуды не пусты
            centroid_value = np.sum(np.array(limited_freqs) * np.array(limited_amps)) / np.sum(limited_amps)
            centroids.append(centroid_value)
        else:
            centroids.append(None)  # Если амплитуды пустые, добавляем None
#------------------------------------002----------
        # Предполагаем, что centroids — это список с вычисленными значениями центроидов
        centroids_sorted = sorted(enumerate(centroids), key=lambda x: x[1], reverse=True)

# Получаем индексы и значения для топ-3, но проверяем, сколько элементов в списке
        top_3_centroids = centroids_sorted[:5]  # Будет содержать максимум 3 элемента

# Формируем строку для обновления окна
        # Формируем строку для обновления окна
        measure_text = "По яркости звука:\n"
        for i, (index, value) in enumerate(top_3_centroids):
            file_name = selected_files[index]  # Получаем имя файла, соответствующее индексу
            measure_text += f"{i+1} место у {file_name}, "


# Обновляем окно с топ-3 значениями центроидов
        window['-MEASURE-'].update(measure_text, text_color='ivory')


#------------------------------------002----------            
            

    # Отрисовка гармоник
        ax.plot(harmonic_freqs, harmonic_amps, color=colors[i])

    # Убираем расширение из имени файла
        base_filename = os.path.splitext(selected_files[i])[0]

    # Отрисовка точек гармоник с обновленным именем файла
    #    ax.scatter(
    #    harmonic_freqs, harmonic_amps,
    #    color=colors[i], marker=markers[i],
    #    label=f'Обертоны: {base_filename}',
    #    linewidth=0.8, s=20
    #)
#-----------------------------------0001-------
        # Установка параметров фона
 
# Отрисовка точек гармоник
        ax.scatter(
            harmonic_freqs, harmonic_amps,
            color=colors[i], marker=markers[i],
            label=f'Обертоны: {base_filename}',
            linewidth=0.9,  # Тонкая линия обводки точки
            s=30            # Уменьшенный размер точек
        )
        # Устанавливаем параметры легенды
        legend = ax.legend(loc='upper right', fontsize=8, bbox_to_anchor=(1.06, 1.05))
        #legend.get_frame().set_facecolor((0, 0.2, 0.4, 0.1))  # Полупрозрачный черный фон
        legend.get_frame().set_alpha(0)  # Полностью прозрачный фон
        legend.get_frame().set_edgecolor('none')  # Убираем границу
        for text in legend.get_texts():
            text.set_color("lightgray")  # Светло-серый текст

#-----------------------------------0001-------
    # Построение линии тренда (полифит первой степени для первых 20 гармоник)
        if len(limited_freqs) > 1:
            p = np.polyfit(limited_freqs, limited_amps, 1)
            ax.plot(limited_freqs, np.polyval(p, limited_freqs), '-', color=colors[i], linewidth=0.85, alpha=0.85)
            
        slopes.append((selected_files[i], p[0]))
        
        # Сортировка по убыванию абсолютного значения наклона
        top_slopes = sorted(slopes, key=lambda x: abs(x[1]))[:3]
        
        
    
# Формирование текста для вывода
        slopes_text = "Линия тренда более пологая - "
        for filepath, slope in top_slopes:
            filename = os.path.splitext(os.path.basename(filepath))[0]
            slopes_text += f"{filename}; затем - "
            
        

# Обновление текста в окне
        window['-SLOPES-'].update(slopes_text, text_color='ivory')

        
#_____график________       
        
    ax.text(
    0.6, 1.0, # Координаты (x, y) в относительных значениях
    f'Сравнение трендов по обертонам', 
    transform=ax.transAxes,  # Привязка координат графика (Axes)
    ha='right',          # Горизонтальное выравнивание текста
    va='bottom',         # Вертикальное выравнивание текста
    fontsize=10, 
    color='lightgray'
)    

# Настройки осей и графика
    ax.grid(True)
    #ax.legend()

# Добавление блоков аннотаций под графиком
    # Настройки позиции для аннотаций
    y_pos = -0.12  # Позиция по Y ниже графика
    x_step = 0.2  # Шаг между аннотациями
    x_offset = 0.08  # Смещение начала вправо

# Добавление блоков аннотаций под графиком
    for i, file_name in enumerate(selected_files[:5]):  # Ограничиваем до 5 файлов
        x_pos = x_offset + i * x_step  # Добавляем смещение x_offset к позиции x_pos
       # Начальная часть текста аннотации
        annotation_text = (
    f"{file_name}\n"
    f"Центроид: {centroids[i]:.0f} Гц\n" if centroids[i] is not None else f"{file_name}\nЦентроид: Нет данных\n"
)


        ax.text(
        x_pos, y_pos, annotation_text,
        ha='center', va='top', transform=ax.transAxes,
        color="ivory", fontsize=8,
        bbox=dict(facecolor="black", edgecolor="black", boxstyle="square,pad=1.5", alpha=0.1)
    )

#_вставить===
    # Сохранение графика в формате PNG через fig
   
    fig.savefig(f"сравнение.png")

    
    

    # Обновление графика в интерфейсе
    figure_canvas_agg.draw()

    return slopes

# --- Функция для отрисовки графика в интерфейсе  =1+2+3=======
def draw_figure(canvas, figure):
    figure_canvas_agg = FigureCanvasTkAgg(figure, canvas)
    figure_canvas_agg.draw()
    figure_canvas_agg.get_tk_widget().pack(side='top', fill='both', expand=1)
    return figure_canvas_agg




# Создание объектов Matplotlib__ГРАФИК============:
fig = mpl_fig.Figure(figsize=(7.6, 4.5), dpi=100)
ax = fig.add_subplot(111)

# Настройка внешней окантовки для всего поля фигуры
fig.patch.set_edgecolor('black')   # Устанавливаем цвет окантовки
fig.patch.set_linewidth(0.8)         # Задаем толщину окантовки


# Установка меток осей внутри графика--x_____y-
ax.xaxis.set_label_coords(0.5, -0.1)
ax.yaxis.set_label_coords(-0.1, 0.5)

# Установка делений на осях внутри графика---работает!!!
ax.tick_params(axis='both', direction='in', length=1, width=1, colors='ivory', labelsize=6)

# Настройка цвета фона графика--
#ax.set_facecolor('#172a3c')
ax.set_facecolor('#0a1929') 
# Настройка цвета сетки
#ax.grid(color='grey', alpha=0.3, linestyle='-')
ax.grid(color='#4f5b66', alpha=0.8, linestyle='-', linewidth=0.7)  
# Настройка цвета поля вокруг графика
fig.set_facecolor('#0a1929')  

# Установка позиции графика без полей вокруг-----работает
ax.set_position([0.05, 0.2, 0.9, 0.77])  # Настройка полей вокруг графика

ax.spines['top'].set_color('none')  # Удаление верхней окантовки
#ax.spines['bottom'].set_color('none')  # Удаление нижней окантовки
ax.spines['left'].set_color('none')  # Удаление левой окантовки
#ax.spines['right'].set_color('none')  # Удаление правой окантовки
#текст по осям внутри графика------
ax.text(0.95, 0.01, 'Frequency, Hz',
       verticalalignment='bottom', horizontalalignment='right',
       transform=ax.transAxes,
        color='ivory', fontsize=6)

ax.text(0.01, 0.8, 'Amplitude, dB', fontsize=6, rotation='vertical', ha='center', va='baseline', color='ivory')

ax.text(
            0.8, 1.0,  # Координаты (x, y) в относительных значениях
            'Визуализация ТЕМБРА медно-духовых инструментов. Тромбон', 
            transform=ax.transAxes,  # Привязка координат графика (Axes)
            ha='right',          # Горизонтальное выравнивание текста
            va='bottom',         # Вертикальное выравнивание текста
            fontsize=10, 
            color='lightgray'
            )
# -------- Конец вашего кода Matplotlib --------


# -----------------------------------Запись-------list---

# Функция измерения уровня шума (Добавлено)
def measure_noise_level():
    """Измеряет базовый уровень шума."""
    p = pyaudio.PyAudio()
    stream = p.open(format=FORMAT, channels=CHANNELS, rate=RATE, input=True, frames_per_buffer=CHUNK)
    print("Измерение уровня шума...")  # Информируем пользователя
    update_status_bar(window, "Измерение уровня шума...", text_color='lime')

    # Сбор данных для анализа шума
    noise_frames = []
    for _ in range(int(RATE / CHUNK * NOISE_DURATION)):
        data = stream.read(CHUNK)
        noise_frames.append(np.frombuffer(data, dtype=np.int16))

    # Закрываем поток
    stream.stop_stream()
    stream.close()
    p.terminate()

    # Рассчитываем среднюю амплитуду шума
    noise_amplitude = np.mean([np.abs(frame).mean() for frame in noise_frames])
    print(f"Средний уровень шума: {noise_amplitude}")
    # Опорное значение для 16-битного аудио
    A_ref = 32767

# Переводим амплитуду в дБ
    noise_db = 20 * np.log10(noise_amplitude / A_ref)
     # Обновление текстового поля статус-бара
    window['-MEASURE-'].update( f"Средний уровень шума: {noise_db:.0f} дБ\nИдет процесс записи", text_color='lime')
    # Обновление текста в окне
    slopes_text = f"Нота: {button_values} , частота = {button_frequency} Гц"
    # Обновление текста в окне
    window['-SLOPES-'].update(slopes_text, text_color='ivory')
        
     # Обновление текстового поля статус-бара
    return noise_amplitude

# Функция записи аудио (Изменена)
def record_audio(selected_file, duration):
    """Записывает аудио в файл с указанным именем и длительностью."""
    global is_recording, frames

    # Сброс данных перед новой записью
    frames = []

    # Измеряем уровень шума (Добавлено)
    noise_level = measure_noise_level()  # Получаем базовый уровень шума
    threshold = noise_level * (10 ** (THRESHOLD_DB / 20))  # Вычисляем порог активации

    # Подготовка потока
    p = pyaudio.PyAudio()
    stream = p.open(format=FORMAT, channels=CHANNELS, rate=RATE, input=True, frames_per_buffer=CHUNK)
    print(f"Порог активации: {threshold:.2f}")

    # Ожидание активации записи (Добавлено)
    print("Ожидание звука для начала записи...")
    update_status_bar(window, "Ожидание звука для начала записи...", text_color='lime')
    start_time = None  # Сбрасываем стартовое время

    # Мониторим входящий сигнал
    while True:
        data = stream.read(CHUNK)
        amplitude = np.abs(np.frombuffer(data, dtype=np.int16)).mean()

        # Проверяем, превышает ли сигнал порог
        if amplitude > threshold:
            print("Запись началась!")  # Информируем пользователя
            update_status_bar(window, "Запись началась!", text_color='red')
            start_time = time.time()  # Фиксируем время старта
            break

    # Основной цикл записи (Существующий код)
    while is_recording and (time.time() - start_time < duration):
        data = stream.read(CHUNK)
        frames.append(data)

    # Закрываем поток (Существующий код)
    stream.stop_stream()
    stream.close()
    p.terminate()
    print("Запись остановлена")
    update_status_bar(window, "Запись остановлена", text_color='red')
    window["-TIEMS-"].update(f"Время записи\n{time.time() - start_time:.1f} s")
    # Сохраняем записанные данные в файл (Существующий код)
    wf = wave.open(selected_file, 'wb')
    wf.setnchannels(CHANNELS)
    wf.setsampwidth(p.get_sample_size(FORMAT))
    wf.setframerate(RATE)
    wf.writeframes(b''.join(frames))
    wf.close()
    print(f"Запись сохранена как {selected_file}")

    # Обновляем интерфейс (Существующий код)
    file_list.append(selected_file)
    window['-FILE-'].update(values=file_list)
    update_status_bar(window, f'Запись сохранена как :\n{selected_file}\nНажми STOP', text_color='lime')
    
    window['-MEASURE-'].update( "Процесс записи завершен.", text_color='red')
    slopes_text = f"Нота: {button_values} , частота = {button_frequency} Гц"
    # Обновление текста в окне
    window['-SLOPES-'].update(slopes_text, text_color='ivory')
        
   
    
# -----------------------------------Запись-------list----    
    
    
# --- Интерфейс PySimpleGUI ---========================
# Определение внешнего вида окна с Canvas    
#фрейм 1--------------------------------------1фрейм
frame1_layout = [
    
    [sg.Listbox(values=file_list, size=(22, 18), key='-FILE-', enable_events=True),
    
     sg.Listbox(values=com_file_list, size=(22, 18), key='-COMFILE-', enable_events=True, select_mode='extended')],
    
    [ sg.InputText(default_text='file_name', size=(16), key='-FILENAME-'), sg.Text('___Record time, s___'), sg.Button('CLEAR', size=(10, 1),key='-CLEAR-') ],

    [sg.StatusBar('Select [note on button]', size=(13), key='-NOTE-'), sg.Slider(range=(5, 10), default_value=5, orientation='h', size=(14, 16),
               key='-DURATION-', resolution=1, tick_interval=1), sg.Button('COMPARE',size=(10, 2),key='-COMPARE-')],

    [sg.Button('RECORD',size=(19, 2),key='-RECORD-',button_color=('white','blue')), sg.Button('STOP',size=(10, 2), key='-STOP-',button_color=('white', 'blue')), sg.Button('ANALYSIS',size=(10, 2), key='-ANALYSIS-')
    
     ],
   [sg.StatusBar("Information:", size=(42, 3), text_color='white',key='-STATUS-')],
   
   
    [ sg.Button('Exit',size=(8, 2), button_color=('white', 'red'),key='-EXIT-'),sg.StatusBar("Tiems:", size=(8, 3), key='-TIEMS-'),sg.StatusBar("Frequency:", size=(8, 3), key='-FREQ-'),sg.StatusBar("Level:", size=(8, 3), key='-LEVEL-')]    
                ]
#фрейм 2---------2---------2----------2------2-----
frame2_layout = [
    [sg.Canvas(size=(760, 410),  background_color='dark blue', key='-CANVAS-')], 

#    [sg.Text('_'  * 80)],

    [sg.Button('Bb2',size=(5, 1),key='button1', bind_return_key=True),
     sg.Button('F3',size=(5, 1),key='button2'), 
     sg.Button('Bb3',size=(5, 1),key='button3'), 
     sg.Button('D4',size=(5, 1),key='button4'), 
     sg.Button('F4',size=(5, 1),key='button5'), 
     sg.Button('Ab4',size=(5, 1),key='button6'), 
     sg.Button('Bb4',size=(5, 1),key='button7'), sg.StatusBar("Measuring:", size=(30, 3), text_color='white',key='-MEASURE-')],

      [sg.Text('Percentile Noise:'),  sg.Text(noise_percentile, key='-NOISE-VALUE-'), 
      sg.Text('   Trimming time, s:'), sg.Text(trim, key='-TRIM-VALUE-'), sg.Push(), sg.StatusBar("Наклон трендов:", size=(55, 2), text_color='white',key='-SLOPES-')
     ], 
     [sg.Slider(range=(75, 95), default_value=noise_percentile, disable_number_display=True, orientation='h', size=(14, 14),
               key='-NOISE-', enable_events=True, resolution=5, tick_interval=5), sg.Text(' '), sg.Slider(range=(0, 5), default_value=1, orientation='h', size=(14, 14), disable_number_display=True,
               key='-TRIM-', enable_events=True, resolution=1, tick_interval=1),
               sg.Push(),  # толкает всё, что дальше, вправо
               sg.Text("Автор: Пиданов, 2025 | Бесплатно. Используйте с упоминанием автора.", size=(70, 1), justification='center', font=('Helvetica', 8), text_color='gray')]
                 ]
 
frame1 = sg.Frame('ListBox File:                             ComFile:', frame1_layout, relief=sg.RELIEF_FLAT)
frame2 = sg.Frame('Graphs', frame2_layout, relief=sg.RELIEF_FLAT)

# Добавление логотипа в интерфейс
#logo = sg.Image(filename=logo_path)

# Создание интерфейса с логотипом
layout = [
    [frame1, frame2]
]
# Создаем окно============================интерфейса++++
window = sg.Window("BrassMouthpieceDemo - Автор: Пиданов, 2025", layout, finalize=True, element_justification='center', resizable=True)



# Пример данных для Aннотаций---------А==
harmonic_frequenciess = [100 * i for i in range(1, 11)]  # Частоты от 100 Гц до 1000 Гц
harmonic_levelss = [-20 + i * 2 for i in range(1, 11)]    # Пример уровней от -20 дБ до 0 дБ

# Позиция для аннотаций
y_pos = -0.12  # Позиция по Y ниже графика
x_step = 0.11  # Шаг по X для каждого блока

# Добавление блоков аннотаций
for i in range(10):
    # Определяем координаты аннотации
    x_pos = i * x_step

    # Текст аннотации
    annotation_text = (f"{harmonic_frequenciess[i]} Hz\n"
                       f"{harmonic_levelss[i]} dB\n"
                       f"  {i+1}")

    # Добавляем аннотацию на график
    ax.text(x_pos, y_pos, annotation_text,
            ha='center', va='top', transform=ax.transAxes,
            color="ivory", fontsize=6, bbox=dict(facecolor="black", edgecolor="black", boxstyle="square,pad=1.5", alpha=0.1))
    
# Добавляем график в окно-------Б====
figure_canvas_agg = draw_figure(window['-CANVAS-'].TKCanvas, fig)

#=draw_figure(window['-CANVAS-'].TKCanvas, fig)

# Обновляем статус-бар после создания окна=-=-=-=-=-
window['-NOTE-'].update(f"Частота = {button_frequency} Hz")
# Функция для обновления состояния в полосе состояния

#def update_status_bar(message):
#    window['-STATUS-'].update(message)
   
# Функция для обновления статус-бара
def update_status_bar(window, message, text_color='lime'):
    window['-STATUS-'].update(message, text_color=text_color)    



    #+конец графика сигнала

# Цикл обработки событий. Ждем событий от пользователя
while True:
    event, values = window.read(timeout=100)
# Если пользователь закрыл окно, завершаем программу и Добавляем таймаут для обновления графического интерфейса
    if event == sg.WIN_CLOSED or event == '-EXIT-':
        break
    #Присваивание кнопкам значение частоты 
    elif event.startswith('button'):
        button_num = int(event[6:]) - 1
        button_frequency = VALUES[button_num]
        diapazon = DIAPAZON[button_num]
        button_values = NOTES[button_num]
        print(f"Для button_num {button_num} - {button_values} из NOTES), diapazon = {diapazon}")
        
        
        window["-NOTE-"].update(f"Частота = {button_frequency} Гц")  # Обновляем значение статус-бара для уровня звука
        slopes_text = f"Нота: {button_values} , частота = {button_frequency} Гц"
    # Обновление текста в окне
        window['-SLOPES-'].update(slopes_text, text_color='ivory')
        
        

    if event == '-TRIM-':
        trim = int(values['-TRIM-'])  # Обновляем значение trim на основе слайдера
        window['-TRIM-VALUE-'].update(trim)  # Обновляем текст для отображения текущего значения trim
        

    if event == '-NOISE-':
        noise_percentile = int(values['-NOISE-']) # Обновляем значение noise_percentile на основе слайдера
        window['-NOISE-VALUE-'].update(noise_percentile)  # Обновляем текст для отображения текущего значения
        print(f"ПРИСВОЕНИЕ значения 4: Button Frequency: {button_frequency}, Trim: {trim}, Noise Percentile: {noise_percentile}")





# Получите значение элемента ввода
    if event == '-RECORD-':
# Получите значение элемента ввода
        is_recording = True
        selected_file = f'{values["-FILENAME-"]}_{button_frequency}Hz.wav' 
# Открываем файл для записи
        duration = int(values['-DURATION-'])  # Время записи в секундах
       
        update_status_bar(window, 'Запись началась...', text_color='lime')
        # Обновление текстового поля статус-бара
        window['-MEASURE-'].update("Запись")
        # Запускаем запись в отдельном потоке, чтобы не блокировать интерфейс
        threading.Thread(target=record_audio, args=(selected_file, duration), daemon=True).start()        
        
        window.refresh()  # Обновляем графический интерфейс
       

    elif event == '-STOP-':   
         # Проверяем, велась ли запись
         # Обновление текстового поля статус-бара
        window['-MEASURE-'].update("СТОП")
        if is_recording:
            is_recording = False  # Останавливаем запись
            sg.popup("Запись завершена!", title="Recording")

            # Добавляем файл в список файлов и обновляем Listbox
            if selected_file:
                file_list.append(selected_file)
                window['-FILE-'].update(values=file_list)

            update_status_bar(window, 'Запись остановлена', text_color='lime')
        else:
            # Если запись не велась, обновляем строку состояния
            update_status_bar(window,'Запись не началась. Файл не добавлен.', text_color='red')
            

#-------------------------------------------
# Обновление визуализации аудио в режиме реального времени========================================
    if is_recording and len(frames) > 0:
        # Преобразование последнего блока данных для отображения в реальном времени:
        audio_chunk = np.frombuffer(frames[-1], dtype=np.int16)
        plot_signal(audio_chunk, figure_canvas_agg, ax)
        # для вывода текущих значений времени частоты и уровня: ---------------

# FFT для частотного спектра
        fft = np.fft.fft(audio_chunk)[:CHUNK // 2]
        freq = np.fft.fftfreq(len(audio_chunk), 1.0 / RATE)[:CHUNK // 2]
#    magnitudes = np.abs(fft)
        magnitudes = 20 * np.log10(np.abs(fft) + 1e-6)  # Амплитуды в dB-------=
 # Обновляем информацию в статус-баре
        rms = 20 * np.log10(np.abs(audio_chunk).mean())
        if np.isnan(rms) or rms == np.inf:
            rms = 0
        window["-LEVEL-"].update(f"RMS Level\n{rms:.0f} dB")
        
        window["-FREQ-"].update(f"Frequency \n{freq[np.argmax(magnitudes)]:.0f} Hz")
        
        

        window.refresh()  # Обновляем графический интерфейс----------------------------------=
              


# Если пользователь выбрал файл в Listbox, выводим его имя в консоль-------------
    if event == '-FILE-':
        print(values['-FILE-'][0])

# Добавляем файл в список выбранных файлов
        com_file_list.append(values['-FILE-'][0])

# Обновляем список выбранных файлов в Listbox
        window['-COMFILE-'].update(values=com_file_list)
      
       
# -----Если пользователь выбрал файл в Listbox compare, выводим его имя в консоль при его наличии---
    if event == '-COMFILE-' and values['-COMFILE-']:
       print(values['-COMFILE-'][0])

# -----Очищаем список com_file_list-----
    elif event == '-CLEAR-':
        if com_file_list and values['-COMFILE-']:
           com_file_list.remove(values['-COMFILE-'][0])
           window['-COMFILE-'].update(values=com_file_list)
           
# --------ВКЛЮЧАЕМ АНАЛИЗ: -----ANALYSIS-----------------
# 
    elif event == '-ANALYSIS-' and not is_recording:
        selected_file = values['-COMFILE-']
        print(f"Файл для анализа: {selected_file}")
        if selected_file:
            try:# Если selected_file является списком, извлекаем первый элемент
                if isinstance(selected_file, list):
                    selected_file = selected_file[0]               
# Шаг 1: Загружаем аудиофайл______________________2_____
                fs, data = wavfile.read(selected_file)  # Считываем данные из выбранного файла
                print(f"Частота дискретизации (fs): {fs}, Длина исходных данных: {len(data)}")                   
# Шаг 2: Проверяем, что данные одномерные (моно), если нет — преобразуем     
                if len(data.shape) > 1:
#                    update_status_bar(window, f"Стерео аудио файл, преобразуем в моно.\nЧастота дискретизации:{fs},\nДлина исходных данных: {len(data)}", text_color='#FFA500')
                    data = data[:, 0] # Берем только 1 канал
#                else:
#                    update_status_bar(window, f"Моно аудио обнаружено, анализируем спектр.\nЧастота дискретизации: {fs},\nДлина исходных данных: {len(data)}", text_color='#00FFFF')
         
                trim = int(values['-TRIM-'])  # Обновляем
                noise_percentile = int(values['-NOISE-'])
                print(f"Значение trim: {trim}, Процент шума: {noise_percentile}") 
                 
#  Начинаем с времени  trim:
                start_sample = int(fs * trim)
# Обрезаем данные начиная с 1-й секунды
                if start_sample < len(data):
                    data = data[start_sample:]
                    print(f"Длина данных после обрезки: {len(data)}")
                else:
                    raise ValueError(f"Файл слишком короткий для начала с {trim}-й секунды.")    
                if len(data) == 0:
                    sg.popup("Ошибка: недостаточно данных после обрезки.")
                else:
# Шаг 3: Вызываем функцию для построения спектра__2______>>>>>>
                    plot_spectrum(data, figure_canvas_agg, ax, button_frequency, trim, noise_percentile, diapazon)
                

            except Exception as e:
                    sg.popup(f"Ошибка при чтении файла: {e}")
        else:
            sg.popup("Пожалуйста, выберите файл из списка -COMFILE- для анализа.")    
# ------------= COMPARE =-----------------------------------

    elif event == '-COMPARE-':
    # Получаем выбранные файлы из Listbox
        selected_files = values['-COMFILE-']
    
        try:
        # Убираем дубликаты и фильтруем только .wav файлы
            selected_files = list(set(file for file in selected_files if file.endswith('.wav')))
        
        # Проверяем, что список файлов содержит от 3 до 5 элементов
            if len(selected_files) < 3:
                sg.popup("Недостаточно файлов .wav для сравнения. Нужно минимум 3 файла.")
            elif len(selected_files) > 5:
                sg.popup("Вы выбрали больше 5 файлов. Пожалуйста, выберите до 5 файлов.")
            else:
            # Список для хранения обработанных данных
                harmonics_data = []
                update_status_bar(window, f"Аудио файлы: {selected_files}")
            
            # Обрабатываем каждый файл
                for file in selected_files:
                    try:
                        harmonic_freqs, harmonic_amps = extract_harmonics(
                        file, trim, diapazon, 40, button_frequency, tolerance_percentage
                    )
                        harmonics_data.append((harmonic_freqs, harmonic_amps))
                    except ValueError as e:
                        sg.popup(f"Ошибка обработки файла {file}: {e}")
       # Вызов функции для отрисовки графика
                slopes = compare_spectrums(harmonics_data, selected_files, figure_canvas_agg, ax)
    
            
            # Дальнейшая обработка данных harmonics_data при необходимости
#                if harmonics_data:
#                    process_harmonics_data(harmonics_data)
        except Exception as e:
            sg.popup(f"Произошла ошибка: {e}")


window.close()